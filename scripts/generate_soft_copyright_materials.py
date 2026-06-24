#!/usr/bin/env python3
"""生成软件著作权申请辅助材料。"""

from __future__ import annotations

import os
import subprocess
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_TIMESTAMP = os.environ.get("SOFT_COPYRIGHT_TIMESTAMP") or datetime.now().strftime("%Y%m%d-%H%M%S")
OUT_DIR = ROOT / "artifacts" / "soft-copyright-application" / OUTPUT_TIMESTAMP
SOFTWARE_NAME = "时安解忧屋综合服务平台"
VERSION = "V1.0.0"
SOURCE_PATHS = [
    "package.json",
    "vite.config.js",
    "index.html",
    "src/**",
    "backend/**",
    "database/schema.sql",
    "scripts/**",
    "configs/release/**",
    "desktop/**",
    "android-shell/**",
    "deploy-h5-to-server.sh",
    "deploy-to-server.sh",
    "deploy-to-staging.sh",
    "rollback-h5-on-server.sh",
    "start-dev.sh",
    "start-h5-preview.sh",
    "start-macos-dev.sh",
    "start-macos.sh",
]


def run_git(args: list[str]) -> str:
    return subprocess.check_output(["git", *args], cwd=ROOT, text=True).strip()


def selected_source_files() -> list[Path]:
    raw = run_git(["ls-files", *SOURCE_PATHS])
    files: list[Path] = []
    excluded_parts = {
        "node_modules",
        "dist",
        "artifacts",
        "backups",
        "server-backup",
        "uploads",
        "__pycache__",
        ".gradle",
        "build",
    }
    excluded_suffixes = {
        ".jpg",
        ".jpeg",
        ".png",
        ".webp",
        ".gif",
        ".svg",
        ".db",
        ".pyc",
        ".map",
        ".icns",
        ".ico",
    }
    excluded_names = {"html2canvas.min.js", "package-lock.json"}
    for line in raw.splitlines():
        path = Path(line)
        if any(part in excluded_parts for part in path.parts):
            continue
        if path.name in excluded_names or path.suffix.lower() in excluded_suffixes:
            continue
        files.append(ROOT / path)
    return files


def collect_source_lines(files: list[Path]) -> list[str]:
    lines: list[str] = []
    for path in files:
        rel = path.relative_to(ROOT)
        try:
            content = path.read_text(encoding="utf-8", errors="ignore").splitlines()
        except OSError:
            continue
        lines.append(f"// ===== 文件：{rel} =====")
        lines.extend(content)
        lines.append("")
    return lines


def generate_source_pdf(lines: list[str], out_file: Path) -> None:
    from reportlab.pdfgen import canvas

    width, height = A4
    left = 16 * mm
    top = height - 16 * mm
    line_height = 4.1 * mm
    lines_per_page = 50
    first = lines[: lines_per_page * 30]
    last = lines[-lines_per_page * 30 :] if len(lines) > lines_per_page * 30 else []
    selected = first + ["// ===== 以下为源程序后 30 页鉴别材料 ====="] + last
    pages = [selected[i : i + lines_per_page] for i in range(0, len(selected), lines_per_page)]

    c = canvas.Canvas(str(out_file), pagesize=A4)
    c.setTitle(f"{SOFTWARE_NAME}{VERSION} 源程序鉴别材料")
    for page_no, page_lines in enumerate(pages[:60], start=1):
        c.setFont("Courier", 6.5)
        c.drawString(left, height - 10 * mm, f"{SOFTWARE_NAME} {VERSION} 源程序鉴别材料 第 {page_no} 页")
        y = top - 8 * mm
        for line_no, line in enumerate(page_lines, start=1):
            safe_line = line.replace("\t", "    ")[:170]
            c.drawString(left, y, f"{line_no:02d} {safe_line}")
            y -= line_height
        c.showPage()
    c.save()


def manual_sections() -> list[tuple[str, list[str]]]:
    return [
        ("一、软件概述", [
            f"{SOFTWARE_NAME}{VERSION} 是面向 H5 和多端构建的综合服务平台。",
            "软件以 uni-app/Vue3 为前端技术基础，以 Python Flask 为后端服务基础，以 SQLite 保存用户、排盘、对话、积分和运营数据。",
            "系统围绕八字、奇门遁甲、紫微斗数、塔罗、择吉、综合 AI 问答、用户账号、会员积分、付费内容、后台管理和多端发行能力等模块提供在线服务。",
            "本说明书用于描述软件的功能构成、运行环境、业务流程、系统模块、数据结构和使用方式。",
        ]),
        ("二、运行环境", [
            "客户端运行环境：支持现代浏览器访问 H5 页面，并可通过 uni-app 构建为微信小程序、支付宝小程序、头条小程序和 App；项目同时提供 Android 壳工程和桌面端壳工程作为多平台发布基础。",
            "前端技术环境：uni-app、Vue3、JavaScript、Pinia、Vite。",
            "服务端技术环境：Python、Flask、SQLAlchemy、SQLite。",
            "桌面端和移动端扩展环境：Electron 桌面端、Android Gradle 工程、HBuilderX/DCloud 打包资源、平台图标和商店材料检查脚本。",
            "第三方服务：DeepSeek/SiliconFlow 兼容 AI 接口，邮件或验证码服务，支付与运营主体可按部署环境接入。",
            "部署方式：前端执行 H5 构建后发布静态资源，后端以 Flask API 提供业务接口，配合脚本完成发布前检查、线上监控、商店材料检查、多端资源打包和回归测试。",
        ]),
        ("三、主要功能", [
            "首页综合 AI：通过自然语言问事引导用户补充问题背景，自动推荐合适术数工具，并生成多术数合参总结。",
            "八字排盘：根据出生日期、时间、性别等信息生成四柱、五行、十神、大运流年等命理信息。",
            "八字专业盘和合盘：提供更细的专业分析视图，以及双方八字合盘参考。",
            "奇门遁甲：支持时家奇门排盘、九宫信息、八门九星、值符值使和问事解读。",
            "紫微斗数：支持紫微斗数命盘生成、宫位展示和运势分析。",
            "塔罗牌：支持塔罗抽牌、牌阵、牌义解释、对话历史和结果保存。",
            "择吉与黄历：提供日期评分、宜忌理由和结合八字的参考建议。",
            "用户中心：提供登录、资料、历史记录、积分中心、会员权益和个人内容管理。",
            "后台管理：提供用户管理、充值确认、积分调整、运营审计、安全检查和基础运维辅助。",
            "多平台发布：提供应用图标校验、商店材料包、桌面端和 Android 壳工程。",
        ]),
        ("四、用户操作流程", [
            "用户进入首页后，可以直接输入具体问题，也可以进入八字、奇门、紫微、塔罗等单项工具。",
            "用户提交问题后，系统根据问题类型判断是否需要出生资料、时间信息、事件背景或抽牌信息。",
            "系统完成排盘或抽牌后，将结构化盘面传入 AI 解读流程，前端以流式方式展示生成结果。",
            "用户可以继续追问，系统保留上下文并将对话保存到历史记录。",
            "涉及积分消耗的功能会在前端展示消耗规则，并在后端进行余额校验、扣减和日志记录。",
            "管理员可以在后台查看用户、充值、积分和操作审计，处理运营异常。",
        ]),
        ("五、系统结构", [
            "前端目录 src/pages 保存主页面和功能页面，src/components 保存公共导航和通用组件。",
            "src/package-tools 保存工具类功能页面，src/package-user 保存用户中心相关页面。",
            "backend 目录保存 Flask API、业务路由、排盘引擎、AI 服务、积分服务和数据模型。",
            "database/schema.sql 保存数据库表结构定义。",
            "scripts 目录保存构建、发布、线上回归、数据库审计、备份、运维、多端打包和商店材料检查脚本。",
            "configs/release 目录保存多平台发行、图标、隐私披露、商店材料和审核账号检查配置。",
            "desktop 目录保存桌面端壳工程入口、预加载脚本、图标和打包配置。",
            "android-shell 目录保存 Android 壳工程、Manifest、Gradle 配置和主 Activity。",
        ]),
        ("六、数据结构", [
            "用户相关表保存账号、登录资料、个人档案和认证状态。",
            "会员积分表保存用户积分余额、会员权益和 AI 次数包。",
            "积分日志表记录积分发放、扣减、退款、充值确认和幂等键。",
            "排盘历史表保存八字、塔罗、奇门、紫微和综合 AI 对话记录。",
            "后台审计表保存管理员操作记录，便于追踪充值确认、积分调整和内容管理行为。",
        ]),
        ("七、安全与运维", [
            "系统对登录态、会员资产、后台接口和积分变动执行权限校验。",
            "涉及写操作的接口采用 CSRF 保护或明确的公开接口边界。",
            "生产数据库默认通过备份、只读审计、恢复演练和线上健康检查进行保护。",
            "发布前通过 preflight、测试、构建、生产监控和线上回归形成验证链路。",
        ]),
        ("八、模块说明", [
            "综合 AI 模块负责问题理解、工具推荐、资料收集、盘面生成、提示词组织和流式总结。",
            "八字模块负责历法换算、四柱计算、五行十神、旺衰格局、大运流年和专业解释。",
            "奇门模块负责起局、九宫排布、符使星门和问事分析。",
            "紫微模块负责命盘生成、宫位信息和星曜分析。",
            "塔罗模块负责牌库校验、抽牌、牌阵解释和历史保存。",
            "积分模块负责余额、扣减、退款、日志、充值套餐和付费内容购买。",
            "运维模块负责健康检查、数据库审计、备份、恢复演练和生产告警。",
        ]),
    ]


def expanded_manual_lines() -> list[tuple[str, str]]:
    rows: list[tuple[str, str]] = []
    for title, paragraphs in manual_sections():
        rows.append(("heading", title))
        for paragraph in paragraphs:
            rows.append(("body", paragraph))
    files = selected_source_files()
    rows.append(("heading", "九、主要源文件清单"))
    for path in files:
        rel = path.relative_to(ROOT)
        suffix = path.suffix.lower()
        if str(rel).startswith("backend/"):
            module_type = "后端服务模块"
            duty = "负责 API 路由、业务处理、数据模型、排盘引擎、AI 调用、积分会员或运营管理等服务端逻辑。"
        elif str(rel).startswith("src/pages") or str(rel).startswith("src/package"):
            module_type = "前端页面模块"
            duty = "负责用户界面展示、表单输入、排盘结果渲染、流式内容展示、历史记录和用户交互。"
        elif str(rel).startswith("src/components"):
            module_type = "前端组件模块"
            duty = "负责导航、通用区块、复用 UI 和跨页面展示逻辑。"
        elif str(rel).startswith("database/"):
            module_type = "数据库结构模块"
            duty = "负责定义用户、会员、积分、排盘记录、对话历史和运营审计等数据表结构。"
        elif str(rel).startswith("scripts/"):
            module_type = "工程脚本模块"
            duty = "负责构建、发布、回归测试、数据库审计、备份恢复、运维检查、多端打包和商店材料校验。"
        elif str(rel).startswith("configs/"):
            module_type = "发行配置模块"
            duty = "负责多端发行范围、应用图标、商店材料、隐私披露、审核账号和发布证据要求配置。"
        elif str(rel).startswith("desktop/"):
            module_type = "桌面端壳工程模块"
            duty = "负责桌面端窗口、预加载桥接、资源图标和桌面应用打包运行能力。"
        elif str(rel).startswith("android-shell/"):
            module_type = "Android 壳工程模块"
            duty = "负责 Android 原生壳、应用清单、入口 Activity、资源图标和移动端打包基础。"
        elif rel.name in {"package.json", "vite.config.js", "index.html"} or rel.suffix == ".sh":
            module_type = "工程入口模块"
            duty = "负责依赖、构建入口、页面入口、部署入口或本地启动流程。"
        else:
            module_type = "项目源程序模块"
            duty = "负责系统运行所需的业务或工程能力。"
        rows.append(("body", f"{rel}：{module_type}，文件类型为 {suffix or '无扩展名'}。"))
        rows.append(("body", f"该文件在系统中的作用：{duty}"))
        rows.append(("body", f"该文件属于{SOFTWARE_NAME}{VERSION} 的源程序组成部分，与其他模块共同完成平台功能。"))
    rows.append(("heading", "十、接口与业务流程说明"))
    route_files = [p for p in files if str(p.relative_to(ROOT)).startswith("backend/") and p.suffix == ".py"]
    for path in route_files:
        rel = path.relative_to(ROOT)
        text = path.read_text(encoding="utf-8", errors="ignore")
        route_count = text.count("@app.route") + text.count("@bp.route")
        if route_count:
            rows.append(("body", f"{rel} 定义约 {route_count} 个接口入口，参与前后端数据交互、用户请求处理和业务状态更新。"))
            rows.append(("body", "接口处理流程通常包括参数读取、登录态或权限校验、业务规则判断、数据库读写、异常处理和 JSON/SSE 响应输出。"))
            rows.append(("body", "涉及积分、会员、充值、后台管理、用户内容和历史记录的接口，需要结合资源归属、幂等键和审计日志保证业务一致性。"))
    rows.append(("heading", "十一、数据库和数据安全说明"))
    rows.extend([
        ("body", "系统数据库以 SQLite 为基础，生产环境通过明确的 DATABASE_URL 指向实际数据库文件。"),
        ("body", "数据库表结构覆盖账号、会员、积分日志、排盘历史、塔罗对话、奇门对话、综合 AI 对话、后台审计和运营记录。"),
        ("body", "积分扣减、充值确认和 AI 生成记录需要保持事务一致性，避免出现余额扣减成功但业务记录缺失，或业务成功但积分未扣减的情况。"),
        ("body", "生产数据库处理遵循先备份、再审计、再验证的原则，不直接以本地库覆盖线上库。"),
        ("body", "审计脚本用于检查关键表、完整性、负积分、孤儿会员记录、孤儿积分日志和重复积分幂等键。"),
    ])
    rows.append(("heading", "十二、发布和维护说明"))
    rows.extend([
        ("body", "发布前执行预检脚本，覆盖依赖、配置、测试、构建和基础安全项。"),
        ("body", "前端 H5 构建完成后发布静态资源，后端 Flask 服务通过部署脚本同步到服务器并重启服务。"),
        ("body", "桌面端和 Android 壳工程通过独立脚本进行资源刷新、安装验证、用户材料包生成和平台状态检查。"),
        ("body", "应用图标和商店材料通过配置化检查脚本固定尺寸、哈希、平台用途和缺口清单。"),
        ("body", "发布后执行生产监控和线上回归，检查首页、八字、奇门、紫微、塔罗、积分中心、登录和关键 API。"),
        ("body", "涉及支付、积分、账号、数据库、上传和后台权限的变更需要额外进行安全复核。"),
        ("body", "备份和恢复演练用于确认生产数据库可恢复，避免上线失败或数据异常时无法回滚。"),
    ])
    rows.append(("heading", "十三、版本和权属说明"))
    rows.extend([
        ("body", f"软件名称：{SOFTWARE_NAME}。"),
        ("body", f"软件版本：{VERSION}。"),
        ("body", "本材料用于个人申请软件著作权登记时整理功能和技术说明，正式申请信息以登记系统填写内容为准。"),
        ("body", "ICP备案主体、支付商户主体和运营主体不等同于软件著作权主体，相关商业运营安排应另行以书面文件确认。"),
    ])
    return rows


def generate_manual_docx(out_file: Path) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"].font.size = Pt(10.5)
    doc.add_heading(f"{SOFTWARE_NAME}{VERSION} 软件设计说明书", level=0)
    for kind, text in expanded_manual_lines():
        if kind == "heading":
            doc.add_heading(text, level=1)
        else:
            doc.add_paragraph(text)
    doc.save(out_file)


def generate_manual_pdf(out_file: Path) -> None:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ChineseTitle",
        parent=styles["Title"],
        fontName="STSong-Light",
        fontSize=20,
        leading=28,
        alignment=1,
    )
    heading_style = ParagraphStyle(
        "ChineseHeading",
        parent=styles["Heading1"],
        fontName="STSong-Light",
        fontSize=14,
        leading=22,
        spaceBefore=10,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "ChineseBody",
        parent=styles["BodyText"],
        fontName="STSong-Light",
        fontSize=10.5,
        leading=18,
        firstLineIndent=21,
        spaceAfter=4,
    )

    story = [Paragraph(f"{SOFTWARE_NAME}{VERSION}<br/>软件设计说明书", title_style), Spacer(1, 14 * mm)]
    row_count = 0
    for kind, text in expanded_manual_lines():
        if kind == "heading":
            story.append(Paragraph(text, heading_style))
        else:
            story.append(Paragraph(text, body_style))
            row_count += 1
        if row_count and row_count % 28 == 0:
            story.append(PageBreak())
    doc = SimpleDocTemplate(
        str(out_file),
        pagesize=A4,
        leftMargin=24 * mm,
        rightMargin=24 * mm,
        topMargin=22 * mm,
        bottomMargin=20 * mm,
        title=f"{SOFTWARE_NAME}{VERSION} 软件设计说明书",
    )
    doc.build(story)


def generate_application_draft(out_file: Path, source_count: int, line_count: int) -> None:
    commit = run_git(["rev-parse", "HEAD"])
    first_commit = run_git(["log", "--reverse", "--format=%h %ad %an %s", "--date=short"]).splitlines()[0]
    out_file.write_text(
        "\n".join([
            f"# {SOFTWARE_NAME}{VERSION} 软著申请字段草稿",
            "",
            f"- 软件全称：{SOFTWARE_NAME}",
            "- 软件简称：时安解忧屋",
            f"- 版本号：{VERSION}",
            "- 著作权人：填写你的个人姓名",
            "- 权利取得方式：原始取得",
            "- 权利范围：全部权利",
            "- 开发方式：独立开发，若有合作事实需按真实情况改为合作开发并补协议",
            "- 软件分类：应用软件 / 行业应用软件 / 互联网服务软件，最终按系统选项选择",
            "- 硬件环境：云服务器、个人电脑、移动终端或浏览器访问设备",
            "- 软件环境：现代浏览器、Python/Flask 服务端、SQLite 数据库、uni-app/Vue3 前端运行环境、Electron 桌面端壳、Android Gradle 壳工程",
            "- 编程语言：JavaScript、Python、SQL、HTML/CSS、Java、Gradle 配置脚本",
            "- 源程序量：以登记系统统计口径填写；本次抽取范围约覆盖 " + str(source_count) + " 个源文件、" + str(line_count) + " 行文本",
            "- 开发完成日期：建议结合真实完成时间填写；当前仓库首个提交记录为 " + first_commit,
            "- 首次发表日期：若已公开上线，填写首次上线日期；若未公开发表，选择未发表或按系统选项填写",
            "- 当前提交哈希：" + commit,
            "",
            "## 软件主要功能",
            "本软件提供玄学综合服务，包含首页综合 AI 问事引导、八字排盘、奇门遁甲、紫微斗数、塔罗牌、择吉黄历、用户账号、历史记录、积分会员、付费内容、后台运营管理、多端发行、应用图标校验、商店材料包、桌面端壳和 Android 壳工程等功能。",
            "",
            "## 备注",
            "ICP备案、微信支付、客服、开票等运营主体信息不等同于软件著作权主体。若公司参与运营，建议另行签署项目权属与运营主体确认文件。",
            "",
        ]),
        encoding="utf-8",
    )


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    files = selected_source_files()
    lines = collect_source_lines(files)
    generate_source_pdf(lines, OUT_DIR / f"{SOFTWARE_NAME}{VERSION}-源程序鉴别材料.pdf")
    generate_manual_docx(OUT_DIR / f"{SOFTWARE_NAME}{VERSION}-软件设计说明书.docx")
    generate_manual_pdf(OUT_DIR / f"{SOFTWARE_NAME}{VERSION}-软件设计说明书.pdf")
    generate_application_draft(OUT_DIR / "软著申请字段草稿.md", len(files), len(lines))
    print(OUT_DIR)
    for path in sorted(OUT_DIR.iterdir()):
        print(path)


if __name__ == "__main__":
    main()
